"""
Document Parser - Extract and parse uploaded documents
"""

import os
import io
from typing import Dict, List, Tuple, Optional
from pathlib import Path


class DocumentParser:
    """
    Parse various document formats (PDF, Word, text, Markdown) and extract content.
    """

    def __init__(self):
        """Initialize the document parser."""
        self.supported_formats = [".pdf", ".docx", ".doc", ".txt", ".md"]

    def parse_file(self, file_path: str) -> Dict[str, any]:
        """
        Parse a document file and extract content.

        Args:
            file_path: Path to the document file

        Returns:
            Dict containing:
                - text: Extracted text content
                - metadata: Document metadata
                - sections: Parsed sections if available
                - format: File format
        """
        file_ext = Path(file_path).suffix.lower()

        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {file_ext}")

        if file_ext == ".pdf":
            return self._parse_pdf(file_path)
        elif file_ext in [".docx", ".doc"]:
            return self._parse_word(file_path)
        elif file_ext in [".txt", ".md"]:
            return self._parse_text(file_path)

        return {"text": "", "metadata": {}, "sections": [], "format": file_ext}

    def _parse_pdf(self, file_path: str) -> Dict[str, any]:
        """Parse PDF file."""
        try:
            import pdfplumber

            content = []
            metadata = {}

            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                metadata["title"] = pdf.metadata.get("Title", "Unknown")

                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)

            return {
                "text": "\n\n".join(content),
                "metadata": metadata,
                "sections": self._extract_sections("\n\n".join(content)),
                "format": ".pdf",
            }
        except ImportError:
            return {
                "text": "PDF parsing requires pdfplumber",
                "metadata": {},
                "sections": [],
                "format": ".pdf",
            }
        except Exception as e:
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "sections": [],
                "format": ".pdf",
            }

    def _parse_word(self, file_path: str) -> Dict[str, any]:
        """Parse Word document."""
        try:
            from docx import Document

            doc = Document(file_path)
            content = []
            metadata = {"paragraphs": len(doc.paragraphs)}

            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)

            text = "\n\n".join(content)
            return {
                "text": text,
                "metadata": metadata,
                "sections": self._extract_sections(text),
                "format": ".docx",
            }
        except ImportError:
            return {
                "text": "Word parsing requires python-docx",
                "metadata": {},
                "sections": [],
                "format": ".docx",
            }
        except Exception as e:
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "sections": [],
                "format": ".docx",
            }

    def _parse_text(self, file_path: str) -> Dict[str, any]:
        """Parse plain text or Markdown file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            file_ext = Path(file_path).suffix.lower()
            return {
                "text": text,
                "metadata": {"lines": len(text.split("\n"))},
                "sections": self._extract_sections(text),
                "format": file_ext,
            }
        except Exception as e:
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "sections": [],
                "format": Path(file_path).suffix.lower(),
            }

    def _extract_sections(self, text: str) -> List[Dict[str, str]]:
        """
        Extract sections from text based on headers.

        Args:
            text: Document text

        Returns:
            List of sections with title and content
        """
        sections = []
        lines = text.split("\n")
        current_section = None
        current_content = []

        for line in lines:
            # Check for markdown headers
            if line.startswith("#"):
                if current_section:
                    sections.append(
                        {"title": current_section, "content": "\n".join(current_content)}
                    )
                current_section = line.lstrip("#").strip()
                current_content = []
            elif line.strip():
                if current_section:
                    current_content.append(line)
                else:
                    current_section = "Introduction"
                    current_content.append(line)

        if current_section and current_content:
            sections.append({"title": current_section, "content": "\n".join(current_content)})

        return sections if sections else [{"title": "Content", "content": text}]

    def parse_text_input(self, text: str) -> Dict[str, any]:
        """
        Parse raw text input.

        Args:
            text: Raw text content

        Returns:
            Parsed content dictionary
        """
        return {
            "text": text,
            "metadata": {"lines": len(text.split("\n")), "words": len(text.split())},
            "sections": self._extract_sections(text),
            "format": ".txt",
        }
