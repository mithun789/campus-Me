"""
Word Document Generator - Generate .docx files with formatting
"""

import io
from typing import Dict, List, Optional
from datetime import datetime
import logging

logger = logging.getLogger(__name__)


class WordGenerator:
    """
    Generate Word documents (.docx) with styles, formatting, and professional layouts.
    """

    def __init__(self):
        """Initialize Word generator."""
        self.font_name = "Calibri"
        self.font_size = 11
        self.line_spacing = 1.5

    def generate_word_doc(
        self,
        title: str,
        content: Dict[str, str],
        author: str = "AI Academic Suite",
        include_toc: bool = True,
        include_citations: bool = False,
        citations: List[str] = None,
    ) -> bytes:
        """
        Generate Word document.

        Args:
            title: Document title
            content: Dictionary of section titles and content
            author: Document author
            include_toc: Include table of contents
            include_citations: Include bibliography
            citations: List of citations

        Returns:
            Word document bytes
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Add title
            title_paragraph = doc.add_paragraph(title)
            title_paragraph.style = "Heading 1"
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add author and date
            metadata = doc.add_paragraph(f"By {author}")
            metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
            metadata = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
            metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            # Add table of contents
            if include_toc:
                toc_paragraph = doc.add_paragraph("Table of Contents")
                toc_paragraph.style = "Heading 2"
                for i, section in enumerate(content.keys(), 1):
                    doc.add_paragraph(f"{i}. {section}", style="List Number")
                doc.add_page_break()

            # Add sections
            for section_title, section_content in content.items():
                section_para = doc.add_paragraph(section_title)
                section_para.style = "Heading 2"

                # Split content into paragraphs
                for para_text in section_content.split("\n\n"):
                    if para_text.strip():
                        p = doc.add_paragraph(para_text)
                        p.paragraph_format.line_spacing = self.line_spacing

            # Add bibliography
            if include_citations and citations:
                doc.add_page_break()
                ref_para = doc.add_paragraph("References")
                ref_para.style = "Heading 2"

                for citation in citations:
                    doc.add_paragraph(citation, style="List Bullet")

            # Save to bytes
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            return doc_buffer.getvalue()

        except ImportError:
            logger.warning("python-docx not available")
            return self._generate_word_fallback(title, content)

    def _generate_word_fallback(self, title: str, content: Dict[str, str]) -> bytes:
        """Fallback Word document generation."""
        try:
            # Create a minimal DOCX-like structure
            import zipfile
            from xml.etree import ElementTree as ET

            docx_content = {
                "[Content_Types].xml": '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                "</Types>",
                "word/document.xml": f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                f"<w:body><w:p><w:r><w:t>{title}</w:t></w:r></w:p>"
                f"{''.join(f'<w:p><w:r><w:t>{sec}: {cnt[:100]}</w:t></w:r></w:p>' for sec, cnt in content.items())}"
                f"</w:body></w:document>",
            }

            # Create DOCX file
            docx_buffer = io.BytesIO()
            with zipfile.ZipFile(docx_buffer, "w") as docx:
                for filename, content_str in docx_content.items():
                    docx.writestr(filename, content_str)

            docx_buffer.seek(0)
            return docx_buffer.getvalue()

        except:
            return b"Word generation failed"

    def add_styles(self, doc_bytes: bytes) -> bytes:
        """
        Add professional styles to Word document.

        Args:
            doc_bytes: Word document bytes

        Returns:
            Styled document bytes
        """
        try:
            from docx import Document
            import io

            doc = Document(io.BytesIO(doc_bytes))

            # Modify existing styles
            styles = doc.styles

            # Update Normal style
            if "Normal" in styles:
                style = styles["Normal"]
                style.font.size = Pt(12)
                style.font.name = "Calibri"

            # Update Heading styles
            for i in range(1, 6):
                heading_name = f"Heading {i}"
                if heading_name in styles:
                    style = styles[heading_name]
                    style.font.size = Pt(14 + (5 - i) * 2)
                    style.font.bold = True

            # Save modified document
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)

            return output_buffer.getvalue()

        except:
            return doc_bytes  # Return original if styling fails

    def extract_text_from_docx(self, doc_bytes: bytes) -> str:
        """
        Extract text from Word document.

        Args:
            doc_bytes: Word document bytes

        Returns:
            Extracted text
        """
        try:
            from docx import Document
            import io

            doc = Document(io.BytesIO(doc_bytes))
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            return text

        except:
            return "Document text extraction failed"
